"""FastAPI server wrapping the multi-agent legal platform.

Usage:
    uv run uvicorn api_server:app --host 0.0.0.0 --port 8000 --reload
"""

import asyncio
import io
import json
import mimetypes
import os
import re
import time
from contextlib import asynccontextmanager
from pathlib import Path

from dotenv import load_dotenv

load_dotenv()

# Agent SDK uses permission_mode="bypassPermissions" in options — no env var needed

_agent_venv = Path(__file__).resolve().parent / ".agent_venv" / "bin"
if _agent_venv.exists():
    os.environ["PATH"] = str(_agent_venv) + os.pathsep + os.environ.get("PATH", "")
    os.environ["VIRTUAL_ENV"] = str(_agent_venv.parent)

from fastapi import FastAPI, UploadFile, File, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel

from api_handlers import (
    AGENTS, SENTINEL, registry, run_agent_turn,
)
from session_manager import Session, OUT_DIR
from profile_manager import (
    get_schema, get_values, save_values, get_status,
    get_all_statuses, reset_profile, extract_placeholders,
)
from skill_runner import list_skills, build_skill_prompt


CORS_ORIGINS = os.getenv("API_CORS_ORIGINS", "http://localhost:35428").split(",")
HEARTBEAT_INTERVAL = 15


@asynccontextmanager
async def lifespan(app: FastAPI):
    eviction_task = asyncio.create_task(_eviction_loop())
    yield
    eviction_task.cancel()


app = FastAPI(title="Legal Agents API", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


async def _eviction_loop():
    while True:
        await asyncio.sleep(300)
        await registry.evict_stale()


# --- Models ---

class CreateSessionRequest(BaseModel):
    name: str | None = None


SESSION_ID_PATTERN = re.compile(r"^[a-f0-9]{12}$")


def _validate_session_id(session_id: str) -> None:
    if not SESSION_ID_PATTERN.match(session_id):
        raise HTTPException(status_code=400, detail="Invalid session ID format")


class ChatRequest(BaseModel):
    session_id: str
    message: str
    accept_edit: bool = True


# --- Endpoints ---

@app.get("/api/health")
async def health():
    return {"status": "ok", "agents": len(AGENTS)}


@app.get("/api/agents")
async def list_agents():
    return [
        {"slug": slug, "description": info["description"]}
        for slug, info in AGENTS.items()
    ]


@app.post("/api/sessions")
async def create_session(body: CreateSessionRequest | None = None):
    session = Session()
    if body and body.name:
        session.name = body.name
    session.save()
    return {
        "session_id": session.id,
        "created_at": session.created_at,
        "sandbox_dir": session.sandbox_dir,
        "output_dir": session.output_dir,
    }


@app.get("/api/sessions")
async def list_sessions():
    return Session.list_sessions()


@app.get("/api/sessions/{session_id}")
async def get_session(session_id: str):
    _validate_session_id(session_id)
    try:
        session = Session.load(session_id)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Session not found")

    charts = []
    out_dir = os.path.join(OUT_DIR, session_id)
    if os.path.isdir(out_dir):
        for fname in sorted(os.listdir(out_dir)):
            if fname.endswith(".chart.json"):
                fpath = os.path.join(out_dir, fname)
                try:
                    with open(fpath, "r") as f:
                        charts.append(json.loads(f.read()))
                except (json.JSONDecodeError, IOError):
                    pass

    return {
        "session_id": session.id,
        "created_at": session.created_at,
        "name": session.name,
        "agents_used": session.agents_used,
        "turns": session.turns,
        "output_dir": session.output_dir,
        "charts": charts,
    }


@app.get("/api/sessions/{session_id}/status")
async def session_status(session_id: str):
    _validate_session_id(session_id)
    async with registry._lock:
        state = registry._sessions.get(session_id)
    if not state:
        return {"session_id": session_id, "is_busy": False}
    is_busy = state.lock.locked()
    return {"session_id": session_id, "is_busy": is_busy}


@app.delete("/api/sessions/{session_id}")
async def delete_session(session_id: str):
    _validate_session_id(session_id)
    try:
        Session.load(session_id)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Session not found")
    Session.delete(session_id)
    return {"status": "deleted", "session_id": session_id}


class RenameRequest(BaseModel):
    name: str


@app.patch("/api/sessions/{session_id}")
async def rename_session(session_id: str, body: RenameRequest):
    _validate_session_id(session_id)
    try:
        session = Session.load(session_id)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Session not found")
    session.rename(body.name)
    return {"session_id": session.id, "name": session.name}


@app.post("/api/sessions/{session_id}/upload")
async def upload_file(session_id: str, file: UploadFile = File(...)):
    _validate_session_id(session_id)
    try:
        state = await registry.load_existing(session_id)
    except FileNotFoundError:
        state = await registry.get_or_create(session_id)

    sandbox = state.session.sandbox_dir
    os.makedirs(sandbox, exist_ok=True)
    dest = os.path.join(sandbox, file.filename)
    content = await file.read()
    with open(dest, "wb") as f:
        f.write(content)

    return {"filename": file.filename, "path": dest, "size": len(content)}


@app.get("/api/sessions/{session_id}/files")
async def list_output_files(session_id: str):
    _validate_session_id(session_id)
    out_dir = os.path.join(OUT_DIR, session_id)
    if not os.path.isdir(out_dir):
        return []
    files = []
    for fname in os.listdir(out_dir):
        if fname.endswith(".chart.json"):
            continue
        fpath = os.path.join(out_dir, fname)
        if os.path.isfile(fpath):
            files.append({
                "filename": fname,
                "size": os.path.getsize(fpath),
                "path": f"sandbox/out/{session_id}/{fname}",
            })
    return files


def _md_to_docx(md_content: str, source_filename: str) -> io.BytesIO:
    """Convert markdown content to a .docx file in memory."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_content.split("\n")
    i = 0
    # Skip YAML frontmatter
    if lines and lines[0].strip() == "---":
        i = 1
        while i < len(lines) and lines[i].strip() != "---":
            i += 1
        i += 1

    in_table = False
    table_rows: list[list[str]] = []

    def _flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        cols = len(table_rows[0])
        tbl = doc.add_table(rows=0, cols=cols)
        tbl.style = "Table Grid"
        for row_idx, row_data in enumerate(table_rows):
            row = tbl.add_row()
            for col_idx, cell_text in enumerate(row_data):
                row.cells[col_idx].text = cell_text.strip()
            if row_idx == 0:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Table detection
        if "|" in stripped and stripped.startswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if cells and all(re.match(r"^[-:]+$", c) for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            _flush_table()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s*", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            _add_formatted_text(p, stripped[2:], default_italic=True)
        elif stripped.startswith("---") or stripped.startswith("***"):
            doc.add_paragraph("_" * 50)
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

        i += 1

    if in_table:
        _flush_table()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_formatted_text(paragraph, text: str, default_italic: bool = False):
    """Parse inline markdown formatting (bold, italic, bold+italic) and add runs."""
    parts = re.split(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|___.*?___|__.*?__|_.*?_)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if default_italic:
                run.italic = True
        elif part.startswith("___") and part.endswith("___"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if default_italic:
                run.italic = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("_") and part.endswith("_") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
            if default_italic:
                run.italic = True


def _strip_md_inline(text: str) -> str:
    """Remove markdown inline formatting markers from text."""
    text = re.sub(r"\*\*\*(.*?)\*\*\*", r"\1", text)
    text = re.sub(r"___(.*?)___", r"\1", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"_(.*?)_", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    return text


def _md_to_xlsx(md_content: str, source_filename: str) -> io.BytesIO:
    """Extract tables from markdown and convert to .xlsx."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    wb.remove(wb.active)

    lines = md_content.split("\n")
    tables: list[tuple[str, list[list[str]]]] = []
    current_heading = source_filename.rsplit(".", 1)[0]
    current_table: list[list[str]] = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("#"):
            heading = re.sub(r"^#+\s*", "", stripped)
            if heading:
                current_heading = heading
        elif "|" in stripped and stripped.startswith("|"):
            cells = [_strip_md_inline(c.strip()) for c in stripped.split("|")[1:-1]]
            if cells and all(re.match(r"^[-:]+$", c) for c in cells):
                continue
            current_table.append(cells)
        else:
            if current_table:
                tables.append((current_heading, current_table))
                current_table = []

    if current_table:
        tables.append((current_heading, current_table))

    if not tables:
        # No tables found — put all content as text in a single sheet
        ws = wb.create_sheet("Content")
        for row_idx, line in enumerate(lines, 1):
            ws.cell(row=row_idx, column=1, value=line)
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf

    header_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    for idx, (title, rows) in enumerate(tables):
        sheet_name = re.sub(r"[\\/*?\[\]:]", "", title)[:31] or f"Table {idx + 1}"
        if sheet_name in wb.sheetnames:
            sheet_name = f"{sheet_name[:27]}_{idx}"
        ws = wb.create_sheet(sheet_name)

        for row_idx, row_data in enumerate(rows, 1):
            for col_idx, cell_value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=cell_value)
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                if row_idx == 1:
                    cell.font = header_font
                    cell.fill = header_fill

        from openpyxl.utils import get_column_letter
        for col_idx in range(1, (len(rows[0]) if rows else 0) + 1):
            max_len = max(
                (len(str(ws.cell(row=r, column=col_idx).value or "")) for r in range(1, len(rows) + 1)),
                default=10,
            )
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 50)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


@app.get("/api/sessions/{session_id}/files/{filename}")
async def download_file(
    session_id: str,
    filename: str,
    format: str = Query(default=None, alias="format"),
):
    _validate_session_id(session_id)
    out_dir = os.path.join(OUT_DIR, session_id)
    fpath = os.path.join(out_dir, filename)
    if not os.path.isfile(fpath):
        raise HTTPException(status_code=404, detail="File not found")

    if format and filename.endswith(".md"):
        with open(fpath, "r", encoding="utf-8") as f:
            md_content = f.read()

        if format == "docx":
            buf = _md_to_docx(md_content, filename)
            dl_name = filename.rsplit(".", 1)[0] + ".docx"
            return StreamingResponse(
                buf,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{dl_name}"'},
            )
        elif format == "xlsx":
            buf = _md_to_xlsx(md_content, filename)
            dl_name = filename.rsplit(".", 1)[0] + ".xlsx"
            return StreamingResponse(
                buf,
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f'attachment; filename="{dl_name}"'},
            )

    media_type = mimetypes.guess_type(fpath)[0] or "application/octet-stream"
    return FileResponse(fpath, filename=filename, media_type=media_type)


@app.post("/api/chat")
async def chat(request: ChatRequest):
    session_id = request.session_id
    _validate_session_id(session_id)

    try:
        state = await registry.load_existing(session_id)
    except FileNotFoundError:
        state = await registry.get_or_create(session_id)

    if state.lock.locked():
        raise HTTPException(status_code=409, detail="Session is busy processing another message")

    async def event_stream():
        queue: asyncio.Queue = asyncio.Queue()

        try:
            await asyncio.wait_for(state.lock.acquire(), timeout=0.5)
        except asyncio.TimeoutError:
            yield _sse_format("error", {"message": "Session is busy processing another message", "type": "BusyError"})
            return

        try:
            agent_task = asyncio.create_task(
                run_agent_turn(state, request.message, queue, request.accept_edit)
            )

            last_heartbeat = time.time()

            while True:
                try:
                    item = await asyncio.wait_for(queue.get(), timeout=0.1)
                except asyncio.TimeoutError:
                    if time.time() - last_heartbeat > HEARTBEAT_INTERVAL:
                        yield _sse_format("heartbeat", {})
                        last_heartbeat = time.time()
                    if agent_task.done():
                        exc = agent_task.exception()
                        if exc:
                            yield _sse_format("error", {"message": str(exc), "type": type(exc).__name__})
                        break
                    continue

                if item is SENTINEL:
                    break

                yield _sse_format(item["event"], item["data"])
                last_heartbeat = time.time()
        finally:
            state.lock.release()

    return StreamingResponse(event_stream(), media_type="text/event-stream", headers={
        "Cache-Control": "no-cache",
        "Connection": "keep-alive",
        "X-Accel-Buffering": "no",
    })


# --- Profile endpoints ---

@app.get("/api/profiles")
async def list_profiles():
    return get_all_statuses()


@app.get("/api/profiles/{slug}")
async def get_profile(slug: str):
    try:
        schema = get_schema(slug)
    except ValueError:
        raise HTTPException(status_code=404, detail="Agent not found")
    return {
        "slug": slug,
        "status": get_status(slug),
        "schema": schema,
        "values": get_values(slug),
    }


class ProfileUpdateRequest(BaseModel):
    values: dict[str, str]


@app.put("/api/profiles/{slug}")
async def update_profile(slug: str, body: ProfileUpdateRequest):
    try:
        result = save_values(slug, body.values)
    except ValueError:
        raise HTTPException(status_code=404, detail="Agent not found")
    return result


@app.post("/api/profiles/{slug}/reset")
async def reset_agent_profile(slug: str):
    try:
        result = reset_profile(slug)
    except ValueError:
        raise HTTPException(status_code=404, detail="Agent not found")
    return result


# --- Skill endpoints ---

@app.get("/api/agents/{slug}/skills")
async def get_agent_skills(slug: str):
    skills = list_skills(slug)
    if skills is None:
        raise HTTPException(status_code=404, detail="Agent not found")
    return skills


class SkillRequest(BaseModel):
    session_id: str
    params: dict[str, str] | None = None


@app.post("/api/agents/{slug}/skills/{skill_name}")
async def run_skill(slug: str, skill_name: str, request: SkillRequest):
    prompt = build_skill_prompt(slug, skill_name, request.params)
    if prompt is None:
        raise HTTPException(status_code=404, detail="Skill not found")

    session_id = request.session_id
    try:
        state = await registry.load_existing(session_id)
    except FileNotFoundError:
        state = await registry.get_or_create(session_id)

    if state.lock.locked():
        raise HTTPException(status_code=409, detail="Session is busy")

    async def event_stream():
        queue: asyncio.Queue = asyncio.Queue()

        try:
            await asyncio.wait_for(state.lock.acquire(), timeout=0.5)
        except asyncio.TimeoutError:
            yield _sse_format("error", {"message": "Session is busy", "type": "BusyError"})
            return

        try:
            agent_task = asyncio.create_task(
                run_agent_turn(state, prompt, queue, True)
            )
            last_heartbeat = time.time()
            while True:
                try:
                    item = await asyncio.wait_for(queue.get(), timeout=0.1)
                except asyncio.TimeoutError:
                    if time.time() - last_heartbeat > HEARTBEAT_INTERVAL:
                        yield _sse_format("heartbeat", {})
                        last_heartbeat = time.time()
                    if agent_task.done():
                        exc = agent_task.exception()
                        if exc:
                            yield _sse_format("error", {"message": str(exc), "type": type(exc).__name__})
                        break
                    continue
                if item is SENTINEL:
                    break
                yield _sse_format(item["event"], item["data"])
                last_heartbeat = time.time()
        finally:
            state.lock.release()

    return StreamingResponse(event_stream(), media_type="text/event-stream", headers={
        "Cache-Control": "no-cache",
        "Connection": "keep-alive",
        "X-Accel-Buffering": "no",
    })


def _sse_format(event: str, data: dict) -> str:
    return f"event: {event}\ndata: {json.dumps(data)}\n\n"
